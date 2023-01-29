from tkinter import *
import yake
import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
from heapq import nlargest
from tkinter import filedialog
from tkinter import messagebox
import docx
from fpdf import FPDF
import numpy as np
import pandas as pd
from tkinter import Canvas
from PIL import ImageTk, Image

global _anslist
_anslist = []

global _summary
_summary = []

class Data:
        def file():
            top = Tk()
            top.geometry("480x180")

            def doc_x():
                doc = docx.Document()

                empty_array = np.array([])

                if(len(_summary) != 0):
                    for y in _summary:
                        doc.add_paragraph(y)

                if(len(_anslist) != 0):
                    for y in _anslist:
                        doc.add_paragraph(y)

                _name = _entry.get()
                _file_name = str(_name+".docx")

                doc.save(_file_name)

                top.destroy()

            def pdf():
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=15)

                empty_array = np.array([])
                if(_anslist == empty_array.size == 0):
                    for x in _summary:
                        pdf.cell(100, 10, txt=x, ln=1, align='L')
                else:
                    for x in _anslist:
                        pdf.cell(100, 10, txt=x, ln=1, align='L')

                _name = _entry.get()
                _file_name = str(_name+".pdf")

                pdf.output(_file_name)

                top.destroy()

            _label_title = Label(top, text="Save_As_File", font=("Arial",15))
            _label_title.place(x=185, y=5)

            _label_entry = Label(top, text="File_Name", font=("Arial",15))
            _label_entry.place(x=10, y=70)

            _entry = Entry(top, font=("Arial",15), width="30")
            _entry.place(x=130, y=70)

            _button_save_as_word = Button(top, text="Save_As_DOCX", font=("Arial",15), fg="blue", command = doc_x)
            _button_save_as_word.place(x=80, y=120)

            _button_save_as_pdf = Button(top, text="Save_As_PDF", font=("Arial",15), fg="red", command = pdf)
            _button_save_as_pdf.place(x=250, y=120)


# Main window
win = Tk()
win.state("zoomed")
win.config(bg="lightblue")
win.title("Extractor")

def csv_file():
    root = Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename()
    
    _df = pd.read_csv(file_path)
    pd.set_option("display.max_columns",None)
    _value = _df.head()

    _textbox.insert(INSERT, _value)
    
def Keyword():
    root = Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename()
    
    file = open(file_path, "r+")
    _read = file.readlines()

    for i in _read:
        kw_extractor = yake.KeywordExtractor(top=10, stopwords=None)
        keywords = kw_extractor.extract_keywords(i)
            
        for kw, v in keywords:
            _textbox.insert(INSERT, kw+"\n")
            _anslist.append(kw)

        
def Save():
    Data.file()

def Summary():
    root = Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename()

    file = open(file_path, "r+")
    _read = file.readlines()

    per = 0.5
    
    for i in _read:
        nlp = spacy.load('en_core_web_sm')
        doc = nlp(i)
        tokens = [token.text for token in doc]
        word_frequencies = {}
        
        for word in doc:
            if word.text.lower() not in list(STOP_WORDS):
                if word.text.lower() not in punctuation:
                    if word.text not in word_frequencies.keys():
                        word_frequencies[word.text] = 1
                    else:
                        word_frequencies[word.text] += 1
                        
        max_frequency = max(word_frequencies.values())
        
        for word in word_frequencies.keys():
            word_frequencies[word] = word_frequencies[word] / max_frequency
            
        sentence_tokens = [sent for sent in doc.sents]
        sentence_scores = {}
        
        for sent in sentence_tokens:
            for word in sent:
                if word.text.lower() in word_frequencies.keys():
                    if sent not in sentence_scores.keys():
                        sentence_scores[sent] = word_frequencies[word.text.lower()]
                    else:
                        sentence_scores[sent] += word_frequencies[word.text.lower()]
                        
        select_length = int(len(sentence_tokens) * per)
        summary = nlargest(select_length, sentence_scores, key=sentence_scores.get)
        final_summary = [word.text for word in summary]
        summary = ''.join(final_summary)

        _summary = summary
        _textbox.insert(INSERT, summary+"\n")

def Clear():
    _textbox.delete("1.0", END)
    _anslist.clear()
    _summary.clear()

def Exit():
    win.destroy()


_label_first = Label(win, text="< Keyword Extractor >", font=("Arial",35), fg="black", bg="lightblue")
_label_first.place(x=530, y=10)

_text = "Use this keyword extraction tool to automatically extract keywords and phrases from all your text data."
_label_second = Label(win, text=_text, font=("Arial",15), fg="black", bg="lightblue")
_label_second.place(x=320, y=100)

_textbox = Text(win, wrap=None, height="27", width="57", font=("Arial",15))
_textbox.place(x=650, y=150)

_button_key = Button(win, text="Extract_Keywords", font=("Arial",15), bg="lime", width="20", command=Keyword)
_button_key.place(x=260, y=350)

_button_sum = Button(win, text="Extract_Summary", font=("Arial",15), bg="lime", width="20", command=Summary)
_button_sum.place(x=260, y=400)

_button_key = Button(win, text="Extract_CSV", font=("Arial",15), width="20", bg="lime", command=csv_file)
_button_key.place(x=260, y=450)

_button_save = Button(win, text="Save_As_File", font=("Arial",15), bg="lime", width="20", command=Save)
_button_save.place(x=260, y=500)

_button_clear = Button(win, text="Clear", font=("Arial",15), width="20", bg="lime", command=Clear)
_button_clear.place(x=260, y=550)

_button_exit = Button(win, text="Exit", font=("Arial",15), width="20", bg="lime", command=Exit)
_button_exit.place(x=260, y=600)
